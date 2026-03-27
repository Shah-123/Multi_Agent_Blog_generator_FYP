"""
export_manager.py — Blog Export Formats (HTML / PDF / DOCX)
============================================================
Converts the final Markdown blog post into multiple publishable formats.

Supported formats:
    - HTML  : Styled, standalone HTML file with embedded CSS
    - PDF   : Professional PDF via weasyprint (requires system GTK/Cairo)
    - DOCX  : Microsoft Word document via python-docx

Usage:
    from Graph.export_manager import export_all
    export_all(markdown_text, "/path/to/output/blog_title", "My Blog Title", ["html", "pdf", "docx"])
"""

import re
import logging
from pathlib import Path
from typing import List, Optional

logger = logging.getLogger("blog_pipeline")


# ============================================================================
# CSS TEMPLATE — Professional blog styling
# ============================================================================

_BLOG_CSS = """
/* AI Content Factory — Blog Export Stylesheet */
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap');

:root {
    --bg-primary: #ffffff;
    --bg-code: #f6f8fa;
    --text-primary: #1a1a2e;
    --text-secondary: #555770;
    --accent: #4361ee;
    --accent-light: #eef2ff;
    --border: #e2e8f0;
    --header-bg: linear-gradient(135deg, #1a1a2e 0%, #16213e 50%, #0f3460 100%);
}

* { margin: 0; padding: 0; box-sizing: border-box; }

body {
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
    color: var(--text-primary);
    background: var(--bg-primary);
    line-height: 1.75;
    font-size: 17px;
    -webkit-font-smoothing: antialiased;
}

.blog-header {
    background: var(--header-bg);
    color: #fff;
    padding: 60px 40px;
    text-align: center;
}

.blog-header h1 {
    font-size: 2.4em;
    font-weight: 700;
    margin-bottom: 12px;
    letter-spacing: -0.5px;
}

.blog-header .meta {
    font-size: 0.95em;
    opacity: 0.8;
}

.blog-content {
    max-width: 780px;
    margin: 0 auto;
    padding: 48px 24px 80px;
}

h1 { font-size: 2.2em; font-weight: 700; margin: 40px 0 20px; color: var(--text-primary); }
h2 { font-size: 1.6em; font-weight: 600; margin: 36px 0 16px; color: var(--text-primary); border-bottom: 2px solid var(--accent); padding-bottom: 8px; }
h3 { font-size: 1.3em; font-weight: 600; margin: 28px 0 12px; color: var(--text-secondary); }
h4 { font-size: 1.1em; font-weight: 600; margin: 20px 0 10px; color: var(--text-secondary); }

p { margin-bottom: 16px; color: var(--text-primary); }

a { color: var(--accent); text-decoration: none; border-bottom: 1px solid transparent; transition: border-color 0.2s; }
a:hover { border-bottom-color: var(--accent); }

strong { font-weight: 600; color: var(--text-primary); }

blockquote {
    border-left: 4px solid var(--accent);
    background: var(--accent-light);
    padding: 16px 20px;
    margin: 20px 0;
    border-radius: 0 8px 8px 0;
    font-style: italic;
    color: var(--text-secondary);
}

code {
    font-family: 'JetBrains Mono', 'Fira Code', monospace;
    background: var(--bg-code);
    padding: 2px 6px;
    border-radius: 4px;
    font-size: 0.88em;
    color: #e83e8c;
}

pre {
    background: #1e1e2e;
    color: #cdd6f4;
    padding: 20px;
    border-radius: 10px;
    overflow-x: auto;
    margin: 20px 0;
    font-size: 0.9em;
    line-height: 1.6;
}
pre code { background: none; color: inherit; padding: 0; }

ul, ol { margin: 12px 0 16px 24px; }
li { margin-bottom: 6px; }

table {
    width: 100%;
    border-collapse: collapse;
    margin: 20px 0;
    font-size: 0.95em;
}
th {
    background: var(--accent);
    color: #fff;
    padding: 12px 16px;
    text-align: left;
    font-weight: 600;
}
td {
    padding: 10px 16px;
    border-bottom: 1px solid var(--border);
}
tr:nth-child(even) { background: #f8f9fa; }

img {
    max-width: 100%;
    height: auto;
    border-radius: 10px;
    margin: 20px 0;
    box-shadow: 0 4px 20px rgba(0,0,0,0.08);
}

hr {
    border: none;
    height: 1px;
    background: var(--border);
    margin: 32px 0;
}

.footer {
    text-align: center;
    padding: 30px;
    color: var(--text-secondary);
    font-size: 0.85em;
    border-top: 1px solid var(--border);
    margin-top: 40px;
}
"""


# ============================================================================
# HTML EXPORT
# ============================================================================

def export_html(markdown_text: str, output_path: str, title: str = "Blog Post") -> Optional[str]:
    """
    Converts Markdown to a styled, standalone HTML file.

    Parameters
    ----------
    markdown_text : str
        The full blog post in Markdown format.
    output_path : str
        Path to write the HTML file (without extension — .html will be appended).
    title : str
        The blog title, shown in the header and <title> tag.

    Returns
    -------
    str or None
        The path to the created HTML file, or None on failure.
    """
    try:
        import markdown as md
    except ImportError:
        logger.warning("⚠️ 'markdown' library not installed. Skipping HTML export. Run: pip install markdown")
        return None

    try:
        # Convert Markdown → HTML body
        extensions = ["fenced_code", "tables", "toc", "nl2br", "sane_lists"]
        html_body = md.markdown(markdown_text, extensions=extensions)

        # Strip the H1 from the body (we render it in the header)
        html_body = re.sub(r'<h1[^>]*>.*?</h1>', '', html_body, count=1, flags=re.DOTALL)

        # Assemble full HTML document
        html_doc = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{_escape_html(title)}</title>
    <style>{_BLOG_CSS}</style>
</head>
<body>
    <header class="blog-header">
        <h1>{_escape_html(title)}</h1>
        <p class="meta">Generated by AI Content Factory</p>
    </header>
    <article class="blog-content">
        {html_body}
    </article>
    <footer class="footer">
        <p>Generated by AI Content Factory &mdash; Multi-Agent Blog Generator</p>
    </footer>
</body>
</html>"""

        out = _ensure_extension(output_path, ".html")
        Path(out).write_text(html_doc, encoding="utf-8")
        logger.info(f"   ✅ Exported HTML: {Path(out).name}")
        return out

    except Exception as e:
        logger.error(f"   ❌ HTML export failed: {e}")
        return None


# ============================================================================
# PDF EXPORT
# ============================================================================

def export_pdf(markdown_text: str, output_path: str, title: str = "Blog Post") -> Optional[str]:
    """
    Converts Markdown to PDF using weasyprint.

    Requires weasyprint + system GTK/Cairo libraries.
    Falls back gracefully if not installed.

    Returns
    -------
    str or None
        The path to the created PDF file, or None on failure.
    """
    try:
        from weasyprint import HTML as WeasyprintHTML
    except (ImportError, OSError) as e:
        logger.warning(
            f"⚠️ Skipping PDF export — weasyprint unavailable: {e}. "
            "PDF requires GTK/Cairo system libs. HTML and DOCX exports still work."
        )
        return None

    try:
        import markdown as md
    except ImportError:
        logger.warning("⚠️ 'markdown' library not installed. Skipping PDF export.")
        return None

    try:
        extensions = ["fenced_code", "tables", "toc", "nl2br", "sane_lists"]
        html_body = md.markdown(markdown_text, extensions=extensions)

        # PDF-optimized HTML (no header banner — just clean content)
        html_doc = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <style>
        {_BLOG_CSS}
        /* PDF overrides */
        body {{ font-size: 14px; }}
        .blog-header {{ padding: 30px 20px; }}
        .blog-content {{ padding: 20px; }}
        @page {{ margin: 2cm; }}
    </style>
</head>
<body>
    <header class="blog-header">
        <h1>{_escape_html(title)}</h1>
        <p class="meta">Generated by AI Content Factory</p>
    </header>
    <article class="blog-content">
        {html_body}
    </article>
</body>
</html>"""

        out = _ensure_extension(output_path, ".pdf")
        WeasyprintHTML(string=html_doc).write_pdf(out)
        logger.info(f"   ✅ Exported PDF: {Path(out).name}")
        return out

    except Exception as e:
        logger.error(f"   ❌ PDF export failed: {e}")
        return None


# ============================================================================
# DOCX EXPORT
# ============================================================================

def export_docx(markdown_text: str, output_path: str, title: str = "Blog Post") -> Optional[str]:
    """
    Converts Markdown to a Word (.docx) document using python-docx.

    Parses headings, paragraphs, bullet lists, bold text, and inline links
    from the Markdown source.

    Returns
    -------
    str or None
        The path to the created DOCX file, or None on failure.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.warning("⚠️ 'python-docx' not installed. Skipping DOCX export. Run: pip install python-docx")
        return None

    try:
        doc = Document()

        # -- Document styles --
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        # -- Title --
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Generated by AI Content Factory")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.font.italic = True

        doc.add_paragraph()  # spacer

        # -- Parse Markdown line by line --
        lines = markdown_text.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Skip empty lines
            if not stripped:
                i += 1
                continue

            # Headings
            if stripped.startswith('#### '):
                doc.add_heading(stripped[5:], level=4)
            elif stripped.startswith('### '):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith('## '):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith('# '):
                doc.add_heading(stripped[2:], level=1)

            # Horizontal rule
            elif stripped in ('---', '***', '___'):
                doc.add_paragraph('_' * 50)

            # Bullet list
            elif stripped.startswith('- ') or stripped.startswith('* '):
                text = stripped[2:]
                para = doc.add_paragraph(style='List Bullet')
                _add_rich_text(para, text)

            # Numbered list
            elif re.match(r'^\d+\.\s', stripped):
                text = re.sub(r'^\d+\.\s', '', stripped)
                para = doc.add_paragraph(style='List Number')
                _add_rich_text(para, text)

            # Blockquote
            elif stripped.startswith('> '):
                text = stripped[2:]
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.5)
                run = para.add_run(text)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x57, 0x70)

            # Code block (skip mermaid/code fences)
            elif stripped.startswith('```'):
                # Collect the code block content
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                if code_lines:
                    code_text = '\n'.join(code_lines)
                    para = doc.add_paragraph()
                    run = para.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

            # Image reference (just add caption text)
            elif stripped.startswith('!['):
                match = re.match(r'!\[([^\]]*)\]\(([^)]*)\)', stripped)
                if match:
                    alt_text = match.group(1)
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run(f"[Image: {alt_text}]")
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

            # Regular paragraph
            else:
                para = doc.add_paragraph()
                _add_rich_text(para, stripped)

            i += 1

        # -- Footer --
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Generated by AI Content Factory — Multi-Agent Blog Generator")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)
        run.font.italic = True

        out = _ensure_extension(output_path, ".docx")
        doc.save(out)
        logger.info(f"   ✅ Exported DOCX: {Path(out).name}")
        return out

    except Exception as e:
        logger.error(f"   ❌ DOCX export failed: {e}")
        return None


# ============================================================================
# CONVENIENCE — Export all requested formats at once
# ============================================================================

def export_all(
    markdown_text: str,
    base_path: str,
    title: str = "Blog Post",
    formats: List[str] = None,
) -> dict:
    """
    Export the blog to multiple formats.

    Parameters
    ----------
    markdown_text : str
        The full blog in Markdown.
    base_path : str
        Base output path WITHOUT extension (e.g., "blogs/my_blog/content/blog_title").
        Each exporter appends its own extension (.html, .pdf, .docx).
    title : str
        Blog title for headers and metadata.
    formats : list of str
        Which formats to export. Subset of ["html", "pdf", "docx"].
        Defaults to ["html"] if None.

    Returns
    -------
    dict
        Mapping of format name → output file path (or None if that format failed).
    """
    if formats is None:
        formats = ["html"]

    results = {}
    format_lower = [f.lower().strip() for f in formats]

    if "html" in format_lower:
        results["html"] = export_html(markdown_text, base_path, title)

    if "pdf" in format_lower:
        results["pdf"] = export_pdf(markdown_text, base_path, title)

    if "docx" in format_lower:
        results["docx"] = export_docx(markdown_text, base_path, title)

    successful = [k for k, v in results.items() if v]
    if successful:
        logger.info(f"   📦 Export complete: {', '.join(f.upper() for f in successful)}")
    else:
        logger.warning("   ⚠️ No exports succeeded.")

    return results


# ============================================================================
# INTERNAL HELPERS
# ============================================================================

def _escape_html(text: str) -> str:
    """Escape HTML special characters."""
    return (
        text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
    )


def _ensure_extension(path: str, ext: str) -> str:
    """Ensure the path ends with the given extension."""
    p = Path(path)
    if p.suffix.lower() != ext.lower():
        return str(p.with_suffix(ext))
    return str(p)


def _add_rich_text(paragraph, text: str):
    """
    Parses simple Markdown inline formatting and adds runs to a python-docx paragraph.
    Supports: **bold**, [link text](url), `inline code`.
    """
    from docx.shared import Pt, RGBColor

    # Pattern to match **bold**, [link](url), or `code`
    pattern = re.compile(
        r'(\*\*(.+?)\*\*)'       # group 1,2: bold
        r'|(\[([^\]]+)\]\(([^)]+)\))'  # group 3,4,5: link
        r'|(`([^`]+)`)'         # group 6,7: inline code
    )

    last_end = 0
    for match in pattern.finditer(text):
        # Add plain text before this match
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])

        if match.group(2):  # Bold
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(4):  # Link — render as "text (url)"
            run = paragraph.add_run(match.group(4))
            run.font.color.rgb = RGBColor(0x43, 0x61, 0xEE)
            run.underline = True
        elif match.group(7):  # Inline code
            run = paragraph.add_run(match.group(7))
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xE8, 0x3E, 0x8C)

        last_end = match.end()

    # Add remaining plain text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])
